import asyncio
from pprint import pprint
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.enum.section import WD_ORIENTATION
from datetime import datetime, timedelta

def create_calculation_doc(
    stone_type="изумрудно-зеленый",
    fraction="70-150 мм",
    dimensions="70*150*20 мм",
    volume="0.86 куб. м",
    weight="1450 кг",
    price_per_kg="38 р.",
    total_price="55 100 р.",
    delivery=True,
    packaging="сетка",
    images=None,
    output_path="расчет.docx"
):
    doc = Document()
    
    # Настройка страницы
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    
    # Заголовок
    heading = doc.add_paragraph()
    heading.style = 'Normal'
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(10)
    heading_run = heading.add_run("Индивидуальный расчет ")
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    
    subheading_run = heading.add_run("под ваш проект")
    subheading_run.font.size = Pt(16)
    
    # Горизонтальная линия
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    
    # Добавляем линию
    p_fmt = p.paragraph_format
    run = p.add_run()
    run.add_tab()
    
    # Создаем таблицу 4x6 (добавляем дополнительную колонку)
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Установка ширины столбцов для альбомной ориентации
    for i, width in enumerate([5, 3, 6, 5, 4, 4.7]):
        for cell in table.columns[i].cells:
            cell.width = Cm(width)
    
    # Первая строка: Эрклез
    cell = table.cell(0, 0)
    cell.text = f"Эрклез\n{stone_type}"
    
    # Объединение ячеек для изображений
    image_cell = table.cell(0, 1)
    image_cell.merge(table.cell(0, 5))
    
    # Вторая строка: Фракция
    table.cell(1, 0).text = "Выбранная вами\nфракция"
    table.cell(1, 1).text = fraction
    
    # Объединение ячеек для примечания
    note_cell = table.cell(1, 2)
    note_cell.merge(table.cell(1, 5))
    note_cell.text = "*Ответственность за выбор размера фракции несет клиент"
    
    # Третья строка: Объем и вес
    volume_cell = table.cell(2, 0)
    volume_cell.text = "Объем и вес по вашим параметрам (д*ш*в)"
    
    # Объединение ячеек для размеров
    dim_cell = table.cell(2, 1)
    dim_cell.merge(table.cell(2, 5))
    dim_cell.text = f"{dimensions} = {volume} = {weight}"
    
    # Четвертая строка: Цена и стоимость
    table.cell(3, 0).text = "Цена\nза кг"
    table.cell(3, 1).text = price_per_kg
    
    # Стоимость без НДС
    table.cell(3, 2).text = f"Стоимость,\nбез НДС"
    table.cell(3, 3).text = total_price
    
    # Добавляем колонку с доставкой
    delivery_cell = table.cell(3, 4)
    
    # Добавляем радиокнопки для доставки
    p = delivery_cell.paragraphs[0]
    p.clear()
    
    # Доставка
    radio_without = "☑" if not delivery else "☐"
    radio_with = "☑" if delivery else "☐"
    p.add_run(f"{radio_without} без учета доставки\n")
    p.add_run(f"{radio_with} с учетом доставки")
    
    # Добавляем колонку с упаковкой
    package_cell = table.cell(3, 5)
    
    # Добавляем радиокнопки для упаковки
    p = package_cell.paragraphs[0]
    p.clear()
    
    # Добавляем Упаковка
    p.add_run("Упаковка\n")
    
    # Сетка с радиокнопкой
    radio = "☑" if packaging == "сетка" else "☐"
    p.add_run(f"{radio} сетка\n")
    
    # Мешки с радиокнопкой
    radio = "☑" if packaging == "мешки" else "☐"
    p.add_run(f"{radio} мешки\n")
    
    # Биг-бэги с радиокнопкой
    radio = "☑" if packaging == "биг-бэги" else "☐"
    p.add_run(f"{radio} биг-бэги")
    
    # Добавление изображений в таблицу
    if images and len(images) > 0:
        p = image_cell.paragraphs[0]
        p.clear()
        for img_path in images:
            try:
                run = p.add_run()
                run.add_picture(img_path, width=Cm(3.5), height=Cm(2.8))
                run.add_text(" ")  # Пробел между изображениями
            except Exception as e:
                print(f"Ошибка добавления изображения {img_path}: {e}")
    
    # Применение форматирования как на первой картинке
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            # Светло-серый фон для всей таблицы
            cell_properties = cell._element.tcPr
            if cell_properties is None:
                cell_properties = OxmlElement('w:tcPr')
                cell._element.append(cell_properties)
            
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F2F2F2')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            cell_properties.append(shading)
            
            # Отступы внутри ячеек
            cell.paragraphs[0].paragraph_format.space_before = Pt(4)
            cell.paragraphs[0].paragraph_format.space_after = Pt(4)
            
            # Размер шрифта
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
    
    # Убираем все границы таблицы
    table.style = 'Table Grid'
    
    # Устанавливаем только нижние границы для всех строк, кроме последней
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            
            # Удаляем все существующие границы
            existing_borders = tcPr.xpath('./w:tcBorders')
            for border in existing_borders:
                tcPr.remove(border)
            
            # Добавляем только нижнюю границу для всех строк, кроме последней
            if i < len(table.rows) - 1:
                tcBorders = OxmlElement('w:tcBorders')
                
                # Создаем пустые границы для всех сторон
                for border_type in ['top', 'start', 'end', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_type}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                
                # Добавляем только нижнюю границу
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '4')  # Толщина линии
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), 'auto')
                tcBorders.append(bottom)
                
                tcPr.append(tcBorders)
            else:
                # Для последней строки - удаляем все границы
                tcBorders = OxmlElement('w:tcBorders')
                for border_type in ['top', 'start', 'end', 'bottom', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_type}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
    
    # Сохранение документа
    doc.save(output_path)
    return output_path

def create_calculation_html(
    stone_name="Эрклез",
    stone_type="изумрудно-зеленый",
    fraction="70-150 мм",
    dimensions="70*150*20 мм", 
    volume="0.86 куб. м", 
    weight="1450 кг",
    price_per_kg="38 р.",
    total_price="55 100 р.", 
    delivery=False,
    packaging="сетка",
    images=None,
    output_path="индивидуальный_расчет.html"
):
    """
    Генерирует HTML-файл с карточкой товара и индивидуальным расчетом
    
    Args:
        stone_name: Название камня
        stone_type: Тип камня
        fraction: Фракция
        dimensions: Размеры (д*ш*в)
        volume: Объем
        weight: Вес
        price_per_kg: Цена за кг
        total_price: Общая стоимость
        delivery: Учитывать доставку (True/False)
        packaging: Тип упаковки (сетка/мешки/биг-бэги)
        images: Словарь с путями к изображениям {'default': 'path', 'dry': 'path', 'wet': 'path', 'lit': 'path'}
               или список путей к изображениям ['path1', 'path2', ...]
        output_path: Путь для сохранения файла
    """
    import os
    from base64 import b64encode
    
    # Встраиваем шрифты в base64
    open_sans_regular_base64 = ""
    open_sans_medium_base64 = ""
    
    # Open Sans Regular для обычного текста
    regular_path = "fonts/OpenSans-Regular.ttf"
    if os.path.exists(regular_path):
        try:
            with open(regular_path, "rb") as font_file:
                font_data = b64encode(font_file.read()).decode('utf-8')
                open_sans_regular_base64 = f"data:font/truetype;charset=utf-8;base64,{font_data}"
                print(f"Regular шрифт загружен: {len(font_data)} символов")
        except Exception as e:
            print(f"Ошибка при чтении шрифта {regular_path}: {e}")
            open_sans_regular_base64 = ""
    else:
        print(f"Файл {regular_path} не найден")
    
    # Open Sans Medium для заголовков
    medium_path = "fonts/open-sans-medium.ttf"
    if os.path.exists(medium_path):
        try:
            with open(medium_path, "rb") as font_file:
                font_data = b64encode(font_file.read()).decode('utf-8')
                open_sans_medium_base64 = f"data:font/truetype;charset=utf-8;base64,{font_data}"
                print(f"Medium шрифт загружен: {len(font_data)} символов")
        except Exception as e:
            print(f"Ошибка при чтении шрифта {medium_path}: {e}")
            open_sans_medium_base64 = ""
    else:
        print(f"Файл {medium_path} не найден")
    
    # Настройка изображений
    image_data = {
        'default': "",
        'dry': "",
        'wet': "",
        'lit': ""
    }
    
    # Проверка и преобразование изображений в base64
    if images:
        if isinstance(images, dict):
            # Если передан словарь с изображениями
            for key, path in images.items():
                if path and os.path.exists(path):
                    try:
                        with open(path, "rb") as img_file:
                            img_data = b64encode(img_file.read()).decode('utf-8')
                            ext = os.path.splitext(path)[1].lower().replace('.', '')
                            if ext not in ['jpg', 'jpeg', 'png', 'gif']:
                                ext = 'jpeg'
                            image_data[key] = f"data:image/{ext};base64,{img_data}"
                    except Exception as e:
                        print(f"Ошибка при чтении изображения {path}: {e}")
        
        elif isinstance(images, list) and len(images) > 0:
            # Если передан список изображений
            keys = list(image_data.keys())
            for i, path in enumerate(images[:4]):  # Берем максимум 4 изображения
                if path and os.path.exists(path):
                    try:
                        with open(path, "rb") as img_file:
                            img_data = b64encode(img_file.read()).decode('utf-8')
                            ext = os.path.splitext(path)[1].lower().replace('.', '')
                            if ext not in ['jpg', 'jpeg', 'png', 'gif']:
                                ext = 'jpeg'
                            image_data[keys[i]] = f"data:image/{ext};base64,{img_data}"
                    except Exception as e:
                        print(f"Ошибка при чтении изображения {path}: {e}")
    
    # Убираем SVG-заглушки - теперь отображаем только существующие изображения
    
    # Настройка упаковки
    packaging_options = {
        'сетка': False,
        'мешки': False,
        'биг-бэги': False
    }
    if packaging in packaging_options:
        packaging_options[packaging] = True
    
    # Генерация ячеек с изображениями только для существующих
    image_cells = ""
    image_labels = {
        'default': '',
        'dry': '',
        'wet': '', 
        'lit': ''
    }
    # image_labels = {
    #     'default': '',
    #     'dry': 'сухой',
    #     'wet': 'влажный', 
    #     'lit': 'подсвеченный'
    # }
    
    for key, label in image_labels.items():
        if image_data[key]:  # Если изображение существует
            if label:  # Если есть подпись
                image_cells += f'<td style="border:none;"><img src="{image_data[key]}" alt="Камень {label}" style="width:130px; height:130px; object-fit:cover;"><br><span style="font-size:14px;">{label}</span></td>'
            else:  # Для default без подписи
                image_cells += f'<td style="border:none;"><img src="{image_data[key]}" alt="Камень" style="width:130px; height:130px; object-fit:cover;"><br></td>'
    
    # HTML шаблон
    html_template = f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Индивидуальный расчет</title>
    <style>
        @page {{
            size: landscape;
            margin: 10mm;
        }}
        @font-face {{
            font-family: 'Open Sans Regular';
            src: url('{open_sans_regular_base64}') format('truetype');
            font-weight: normal;
            font-style: normal;
        }}
        @font-face {{
            font-family: 'Open Sans Medium';
            src: url('{open_sans_medium_base64}') format('truetype');
            font-weight: normal;
            font-style: normal;
        }}
        body {{  
            margin: 0;
            padding: 0;
            height: 100%;
            margin-left: 45px;
            margin-right: 45px;
            font-family: 'Open Sans Regular', Arial, sans-serif;
        }}
        .card {{
            width: 100%;
            max-width: 100%;
            box-sizing: border-box;
            padding: 15px;
            display: flex;
            flex-direction: column;
        }}
        h1 {{
            font-size: 30px;
            margin-bottom: 15px;
            font-family: 'Open Sans Medium', Arial, sans-serif;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 15px;
        }}
        td {{
            padding: 10px;
            border-top: 1px solid #333;
            border-bottom: 1px solid #333;
            vertical-align: middle;
        }}
        /* Убираем все границы у последней строки */
        tr:last-child td {{
            border: none;
        }}
        .stone-name {{
            font-size: 20px;
            font-family: 'Open Sans Medium', Arial, sans-serif;
        }}
        .stone-color {{
            font-size: 22px;
            font-family: 'Open Sans Medium', Arial, sans-serif;
        }}
        .stone-images {{
            display: flex;
            justify-content: space-between;
        }}
        .stone-image {{
            text-align: center;
        }}
        .stone-image img {{
            width: 130px;
            height: 130px;
            object-fit: cover;
        }}
        .stone-image p {{
            margin-top: 5px;
            font-size: 14px;
        }}
        .fraction {{
            font-size: 20px;
            font-family: 'Open Sans Medium', Arial, sans-serif;
        }}
        .price {{
            font-size: 18px;
        }}
        .checkbox-container {{
            display: flex;
            align-items: center;
            margin-bottom: 8px;
            font-size: 20px;
        }}
        .checkbox-container input[type="radio"] {{
            margin-right: 8px;
        }}
        .checkbox-circle {{
            display: inline-block;
            width: 18px;
            height: 18px;
            border: 2px solid #333;
            border-radius: 50%;
            margin-right: 8px;
            position: relative;
            background: #fff;
        }}
        .checkbox-circle.checked::after {{
            content: '';
            position: absolute;
            top: 2px;
            left: 2px;
            width: 14px;
            height: 14px;
            background: url('data:image/svg+xml;utf8,<svg width="14" height="14" viewBox="0 0 16 16" fill="none" xmlns="http://www.w3.org/2000/svg"><path d="M4 8.5L7 11.5L12 5.5" stroke="%23000" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"/></svg>') center center no-repeat;
        }}
        /* Вертикальная граница после первой колонки только для первых трёх строк */
        .main-table tr:nth-child(-n+3) td:first-child {{
            border-right: 1.5px solid #333;
        }}
    </style>
</head>
<body>
    <div class="card">
        <h1>Индивидуальный расчет <span style="font-weight: normal;">под ваш проект</span></h1>
        
        <table class="main-table">
            <tr>
                <td style="width: 25%;">
                    <div class="stone-name">Эрклез</div>
                    <div class="stone-name">{stone_name}</div>
                    <div class="stone-color">{stone_type}</div>
                </td>
                <td>
                    <table style="width:100%; border:none; border-collapse:collapse; text-align:center;">
                        <tr>
                            {image_cells}
                        </tr>
                    </table>
                </td>
            </tr>
            
            <tr>
                <td>
                    <div>Выбранная вами</div>
                    <div class="fraction">фракция</div>
                </td>
                <td>
                    <div style="font-size: 22px; display: inline-block; width: 35%;">{fraction}</div>
                    <div style="text-align: right; display: inline-block; width: 55%; float: right;">
                        *Ответственность за выбор<br>
                        размера фракции несет клиент
                    </div>
                </td>
            </tr>
            
            <tr>
                <td>
                    <div><strong>Объем и вес</strong> по вашим</div>
                    <div>параметрам (д*ш*в)</div>
                </td>
                <td>
                    
                    
                    <div style="font-size: 20px;">{dimensions} {volume} {weight}</div>
                </td>
            </tr>
            
            <tr>
                <td style="vertical-align:top; padding-top: 4.8%;">
                    <div class="price">Цена<br>за кг</div>
                    <div style="font-size: 22px; font-weight: bold;">{price_per_kg}</div>
                </td>
                <td colspan="1" style="padding:0;" colspan="2">
                    <table style="width:100%; border:none; border-collapse:collapse;">
                        <tr>
                            <td style="vertical-align:top; border:none; padding:6% 20px 0 0; width:35%;">
                                <div class="price">Стоимость,<br>без НДС</div>
                                <div style="font-size: 22px; font-weight: bold;">{total_price}</div>
                            </td>
                            <td style="vertical-align:top; border:none; padding:7% 20px 0 0; width:40%;">
                                <div class="checkbox-container">
                                    <span class="checkbox-circle {"checked" if not delivery else ""}"></span>
                                    <span style="position: relative; top: -6px;">без учета доставки</span>
                                </div>
                                <div class="checkbox-container">
                                    <span class="checkbox-circle {"checked" if delivery else ""}"></span>
                                    <span style="position: relative; top: -6px;">с учетом доставки</span>
                                </div>
                            </td>
                            <td style="vertical-align:top; border:none; width:40%;">
                                <div class="price">Упаковка</div>
                                <div style="margin-top: 10px;">
                                    <div class="checkbox-container">
                                        <span class="checkbox-circle {"checked" if packaging_options['сетка'] else ""}"></span>
                                        <span style="position: relative; top: -6px;">сетка</span>
                                    </div>
                                    <div class="checkbox-container">
                                        <span class="checkbox-circle {"checked" if packaging_options['мешки'] else ""}"></span>
                                        <span style="position: relative; top: -6px;">мешки</span>
                                    </div>
                                    <div class="checkbox-container">
                                        <span class="checkbox-circle {"checked" if packaging_options['биг-бэги'] else ""}"></span>
                                        <span style="position: relative; top: -5px;">биг-бэги</span>
                                    </div>
                                </div>
                            </td>
                        </tr>
                    </table>
                </td>
            </tr>
        </table>
    </div>
</body>
</html>"""
    
    # Сохранение HTML файла
    with open(output_path, 'w', encoding='utf-8') as html_file:
        html_file.write(html_template)
    
    return output_path

def replace_date_placeholder(pdf_path, output_path=None):
    """
    Заменяет текст "{дата+3}" на "21.03.25" в PDF файле
    
    Args:
        pdf_path: Путь к PDF файлу
        output_path: Путь для сохранения измененного файла (по умолчанию - pdf_path)
    """
    # Сначала пробуем использовать PyMuPDF для реального редактирования
    try:
        import fitz  # PyMuPDF
        
        print("Используем PyMuPDF для редактирования PDF...")
        doc = fitz.open(pdf_path)
        # Подключаем пользовательский шрифт Open Sans Medium, если доступен
        try:
            import os
            font_path = os.path.join(os.path.dirname(__file__), "fonts", "open-sans-medium.ttf")
            custom_fontname = None
            if os.path.exists(font_path):
                try:
                    # Регистрируем шрифт под явным именем, чтобы гарантированно применился
                    explicit_fontname = "opensans_medium"
                    doc.insert_font(fontname=explicit_fontname, fontfile=font_path)
                    custom_fontname = explicit_fontname
                    print(f"Подключен шрифт: {font_path}")
                except Exception as font_err:
                    print(f"Не удалось подключить шрифт '{font_path}': {font_err}")
            else:
                print(f"Шрифт не найден: {font_path}, будет использован шрифт документа")
        except Exception as path_err:
            print(f"Ошибка подготовки пути к шрифту: {path_err}")
            custom_fontname = None
        
        replacements_count = 0
        for page_num in range(len(doc)):
            page = doc[page_num]
            text_instances = page.search_for("{дата+3}")
            
            if text_instances:
                print(f"Найдена дата на странице {page_num + 1}, выполняем замену...")
                for inst in text_instances:
                    # Получаем информацию о текущем тексте для сохранения форматирования
                    try:
                        # Анализируем текущий текст для определения форматирования
                        text_dict = page.get_text("dict")
                        font_info = None
                        
                        # Ищем информацию о шрифте для данного текста
                        for block in text_dict["blocks"]:
                            if "lines" in block:
                                for line in block["lines"]:
                                    for span in line["spans"]:
                                        if "{дата+3}" in span["text"]:
                                            font_info = {
                                                "font": span.get("font", "helv"),
                                                "size": span.get("size", 24),
                                                "color": span.get("color", 0),
                                                "flags": span.get("flags", 0)
                                            }
                                            break
                                    if font_info:
                                        break
                            if font_info:
                                break
                        
                        # Если не удалось получить информацию о шрифте, используем значения по умолчанию
                        if not font_info:
                            font_info = {"font": "helv", "size": 24, "color": 0, "flags": 0}
                        
                        print(f"Используем форматирование: шрифт={font_info['font']}, размер={font_info['size']}")
                        
                        # Заменяем текст с сохранением форматирования
                        page.add_redact_annot(inst, fill=(1, 1, 1))  # Белый фон
                        page.apply_redactions()
                        
                        # Добавляем новый текст с оригинальным форматированием
                        # Используем точную позицию и правильные параметры
                        # inst.tl - top-left, inst.bl - bottom-left для лучшего позиционирования
                        # Опускаем текст на 20pt вниз относительно исходной позиции
                        text_pos = (inst.x0, inst.y0 + 25)
                        nextDate = datetime.now()+ timedelta(days=4)
                        print(nextDate.strftime("%d.%m.%y"))
                        # Нормализуем цвет: PyMuPDF ожидает кортеж RGB (0..1)
                        color_value = font_info.get("color", 0)
                        if not isinstance(color_value, (tuple, list)):
                            color_value = (0, 0, 0)
                        page.insert_text(
                            text_pos,
                            nextDate.strftime("%d.%m.%y") ,
                            fontname=(custom_fontname or font_info["font"]),
                            fontsize=font_info["size"],
                            color=color_value
                        )
                        
                        replacements_count += 1
                        print(f"Замена {replacements_count} выполнена на странице {page_num + 1} с сохранением форматирования")
                        
                    except Exception as format_error:
                        print(f"Ошибка при сохранении форматирования: {format_error}")
                        # Fallback: простая замена без сохранения форматирования
                        page.add_redact_annot(inst, fill=(1, 1, 1))
                        page.apply_redactions()
                        # Опускаем текст на 20pt вниз относительно исходной позиции
                        text_pos = (inst.x0, inst.y0 + 25)
                        page.insert_text(
                            text_pos,
                            nextDate.strftime("%d.%m.%y"),
                            fontname=(custom_fontname or "helv"),
                            fontsize=24,
                            color=(0, 0, 0)
                        )
                        replacements_count += 1
                        print(f"Замена {replacements_count} выполнена на странице {page_num + 1} (без форматирования)") 
        
        if output_path is None:
            output_path = pdf_path.replace('.pdf', '_updated.pdf')
            
        # Сохраняем в новый файл, чтобы избежать ошибки "save to original must be incremental"
        doc.save(output_path, incremental=False)
        doc.close()
        
        print(f"PDF файл сохранен с заменой: {output_path}")
        print(f"Всего замен: {replacements_count}")
        
        # Проверяем результат
        check_doc = fitz.open(output_path)
        still_present = 0
        for check_page_num in range(len(check_doc)):
            check_text = check_doc[check_page_num].get_text()
            if "{дата+3}" in check_text:
                still_present += 1
                print(f"ВНИМАНИЕ: Текст '{{дата+3}}' все еще присутствует на странице {check_page_num + 1}")
        
        check_doc.close()
        
        if still_present == 0:
            print("✅ Замена успешно выполнена во всех местах!")
        else:
            print(f"⚠️ Замена выполнена частично: {still_present} страниц содержат старый текст")
        
        return output_path
        
    except ImportError:
        print("PyMuPDF не установлен, используем pypdf...")
    except Exception as e:
        print(f"Ошибка при использовании PyMuPDF: {e}")
        print("Переключаемся на pypdf...")
    
    # Fallback на pypdf (только для чтения и проверки)
    try:
        from pypdf import PdfReader, PdfWriter
        
        print("Используем pypdf для проверки PDF...")
        reader = PdfReader(pdf_path)
        
        # Проверяем наличие текста для замены
        found_pages = []
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if "{дата+3}" in page_text:
                found_pages.append(page_num + 1)
                print(f"Найдена дата на странице {page_num + 1}")
        
        if found_pages:
            print(f"Текст '{{дата+3}}' найден на страницах: {found_pages}")
            print("⚠️ Замена не может быть выполнена через pypdf")
            print("💡 Рекомендуется установить PyMuPDF: pip install PyMuPDF")
        else:
            print("Текст '{дата+3}' не найден в PDF")
        
        # Создаем копию файла без изменений
        if output_path is None:
            output_path = pdf_path
            
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
            
        with open(output_path, "wb") as output_file:
            writer.write(output_file)
        
        print(f"PDF файл сохранен без изменений: {output_path}")
        return output_path
        
    except Exception as e:
        print(f"Критическая ошибка при работе с PDF: {e}")
        return None

async def main(dealID):
    from workBitrix import get_all_info
    frakcia,ypakovka,dostavka,opportunity,productName,images,productPrice,obem_po_porametram= await get_all_info(dealID)
    
#     create_calculation_doc(
#         stone_type="изумрудно-зеленый",
#         fraction="70-150 мм",
#         dimensions="70*150*20 мм",
#         volume="0.86 куб. м",
#         weight="1450 кг",
#         price_per_kg="38 р.",
#         total_price="55 100 р.",
#         delivery=True,
#         packaging="сетка",
#         images=None,
#     output_path="индивидуальный_расчет.docx"
# )

    # Генерация HTML-карточки с использованием скриншотов в качестве изображений
    # images_list = ["Снимок экрана 2025-04-07 в 14.49.01.png", "Снимок экрана 2025-04-07 в 14.49.15.png"]
    create_calculation_html(
        stone_name=productName,
        # stone_type="изумрудно-зеленый",
        stone_type="",
        fraction=frakcia,
        dimensions=obem_po_porametram,
        volume='',
        weight="", 
        price_per_kg=productPrice, 
        total_price=opportunity, 
        delivery=dostavka, 
        packaging=ypakovka,
        images=images,
        output_path="индивидуальный_расчет.html"
    )
    from insert_html_to_pdf import insert_html_page_to_pdf
    
    insert_html_page_to_pdf(
        html_path="индивидуальный_расчет.html",
        # base_pdf_path="кп 3 вариант.pdf",
        base_pdf_path="кп 3 вариант removed.pdf",
        output_pdf_path="кп 3 вариант_с_вставкой.pdf",
        # insert_after_page=1
    )
    
    # Заменяем "{дата+3}" на "21.03.25" в созданном PDF
    updated_pdf_path = replace_date_placeholder("кп 3 вариант_с_вставкой.pdf")
    if updated_pdf_path and updated_pdf_path != "кп 3 вариант_с_вставкой.pdf":
        print(f"PDF с заменой даты сохранен в: {updated_pdf_path}")
    else:
        print("Замена даты не выполнена")
    
    # from workBitrix import upload_file_to_deal
    # await upload_file_to_deal(dealID, "кп 3 вариант_с_вставкой_updated.pdf")

    import os
    pprint(images)
    for key, value in images.items():
        try:
            os.remove(value)
        except Exception as e:
            continue

if __name__ == "__main__":
    # asyncio.run(main(8342))
    asyncio.run(main(8442))
